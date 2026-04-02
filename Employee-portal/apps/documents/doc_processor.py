"""
Utilities for document processing and extraction
"""
import os
import logging
from pathlib import Path
from typing import List, Tuple
import PyPDF2
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Process documents and extract text for vector embedding
    """
    
    ALLOWED_TYPES = ['pdf', 'docx', 'doc', 'txt']
    
    @staticmethod
    def extract_text(file_path: str, file_type: str) -> Tuple[str, int]:
        """
        Extract text from document
        
        Returns:
            Tuple of (text, word_count)
        """
        try:
            if file_type.lower() == 'pdf':
                return DocumentProcessor._extract_pdf(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                return DocumentProcessor._extract_docx(file_path)
            elif file_type.lower() == 'txt':
                return DocumentProcessor._extract_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def _extract_pdf(file_path: str) -> Tuple[str, int]:
        """Extract text from PDF"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
        except Exception as e:
            logger.error(f"Error reading PDF: {str(e)}")
            raise
        
        word_count = len(text.split())
        return text, word_count
    
    @staticmethod
    def _extract_docx(file_path: str) -> Tuple[str, int]:
        """Extract text from DOCX"""
        text = ""
        try:
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
        except Exception as e:
            logger.error(f"Error reading DOCX: {str(e)}")
            raise
        
        word_count = len(text.split())
        return text, word_count
    
    @staticmethod
    def _extract_txt(file_path: str) -> Tuple[str, int]:
        """Extract text from TXT"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except Exception as e:
            logger.error(f"Error reading TXT: {str(e)}")
            raise
        
        word_count = len(text.split())
        return text, word_count
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
        """
        Split text into chunks for embedding
        
        Args:
            text: Full text to chunk
            chunk_size: Characters per chunk
            chunk_overlap: Overlap between chunks
        
        Returns:
            List of text chunks
        """
        chunks = []
        sentences = text.split('.')
        
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip() + "."
            
            if len(current_chunk) + len(sentence) <= chunk_size:
                current_chunk += " " + sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # Handle overlap
                overlap_text = current_chunk[-chunk_overlap:] if len(current_chunk) > chunk_overlap else current_chunk
                current_chunk = overlap_text + " " + sentence
        
        # Add last chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return [chunk for chunk in chunks if chunk]  # Remove empty chunks
    
    @staticmethod
    def get_file_type(file_name: str) -> str:
        """Get file type from filename"""
        ext = Path(file_name).suffix.lower().strip('.')
        return ext if ext in DocumentProcessor.ALLOWED_TYPES else None


class TextProcessor:
    """
    Process and clean text for better embedding
    """
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize text"""
        import re
        
        # Remove extra whitespace
        text = ' '.join(text.split())
        
        # Remove special characters but keep some punctuation
        text = re.sub(r'[^\w\s\.\,\:\;\!\?\-\(\)]', '', text)
        
        return text.strip()
    
    @staticmethod
    def extract_keywords(text: str, num_keywords: int = 10) -> List[str]:
        """Extract keywords using simple frequency analysis"""
        import re
        from collections import Counter
        
        # Simple word frequency
        words = re.findall(r'\w+', text.lower())
        
        # Filter common words
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'do'
        }
        
        filtered_words = [w for w in words if w not in stop_words and len(w) > 3]
        
        most_common = Counter(filtered_words).most_common(num_keywords)
        
        return [word for word, _ in most_common]
